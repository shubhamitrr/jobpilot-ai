"""
Extracts plain text from uploaded resume files (PDF / DOCX).
Includes basic file-type/size validation to guard against malicious uploads.
"""
import os
from fastapi import UploadFile, HTTPException

from pypdf import PdfReader
from docx import Document

from app.config import settings

ALLOWED_EXTENSIONS = {".pdf", ".docx"}


def validate_upload(file: UploadFile, content: bytes) -> str:
    """Validates extension and size. Returns normalized extension (e.g. '.pdf')."""
    ext = os.path.splitext(file.filename or "")[1].lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=400, detail="Only PDF and DOCX files are supported.")

    max_bytes = settings.MAX_UPLOAD_MB * 1024 * 1024
    if len(content) > max_bytes:
        raise HTTPException(status_code=400, detail=f"File too large. Max {settings.MAX_UPLOAD_MB}MB.")

    if len(content) == 0:
        raise HTTPException(status_code=400, detail="Uploaded file is empty.")

    return ext


def extract_text(filepath: str, ext: str) -> str:
    """Extracts raw text from a resume file on disk."""
    try:
        if ext == ".pdf":
            return _extract_pdf(filepath)
        elif ext == ".docx":
            return _extract_docx(filepath)
        else:
            raise HTTPException(status_code=400, detail="Unsupported file type.")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=422, detail=f"Failed to parse resume file: {e}")


def _extract_pdf(filepath: str) -> str:
    reader = PdfReader(filepath)
    text_parts = []
    for page in reader.pages:
        text = page.extract_text() or ""
        text_parts.append(text)
    full_text = "\n".join(text_parts).strip()
    if not full_text:
        raise HTTPException(
            status_code=422,
            detail="Could not extract text from this PDF (it may be a scanned image). "
                   "Please upload a text-based PDF or a DOCX file.",
        )
    return full_text


def _extract_docx(filepath: str) -> str:
    doc = Document(filepath)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # also pull text from tables (skills tables etc.)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    full_text = "\n".join(paragraphs).strip()
    if not full_text:
        raise HTTPException(status_code=422, detail="This DOCX file appears to be empty.")
    return full_text
